"""Professional ATS-safe resume rendering. Distinct layouts, no internal notes on exports."""

from __future__ import annotations

import io
from typing import Any

from fpdf import FPDF


TEMPLATES = {
    "ats_classic": "ATS Classic",
    "modern_ats": "Modern ATS",
    "technical": "Technical",
    "graduate": "Graduate",
    "executive": "Executive",
    "compact": "Compact",
    "portfolio": "Portfolio",
    "two_tone": "Two-tone",
}

TEMPLATE_LAYOUT = {
    "ats_classic": ["summary", "skills", "experience", "projects", "education"],
    "modern_ats": ["summary", "skills", "experience", "education", "projects"],
    "technical": ["skills", "projects", "experience", "education", "summary"],
    "graduate": ["education", "projects", "skills", "experience", "summary"],
    "executive": ["summary", "experience", "education", "skills", "projects"],
    "compact": ["summary", "experience", "skills", "education", "projects"],
    "portfolio": ["summary", "projects", "skills", "experience", "education"],
    "two_tone": ["summary", "experience", "projects", "skills", "education"],
}

TEMPLATE_COLOR = {
    "ats_classic": (28, 28, 28),
    "modern_ats": (176, 82, 38),
    "technical": (18, 78, 102),
    "graduate": (36, 92, 72),
    "executive": (28, 32, 48),
    "compact": (55, 55, 55),
    "portfolio": (132, 52, 38),
    "two_tone": (22, 58, 88),
}


class ResumePDF(FPDF):
    def __init__(self, accent: tuple[int, int, int], sidebar: bool = False):
        super().__init__(format="Letter")
        self.accent = accent
        self.sidebar = sidebar
        self.set_auto_page_break(auto=True, margin=16)

    def header(self):
        if self.sidebar:
            self.set_fill_color(*self.accent)
            self.rect(0, 0, 8, self.h, "F")

    def footer(self):
        self.set_y(-12)
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(140, 140, 140)
        self.cell(0, 8, str(self.page_no()), align="C")


def _txt(value: Any) -> str:
    raw = str(value or "").replace("\r", "").replace("\t", " ")
    raw = raw.encode("latin-1", "replace").decode("latin-1")
    parts: list[str] = []
    for word in raw.split(" "):
        while len(word) > 70:
            parts.append(word[:70])
            word = word[70:]
        parts.append(word)
    return " ".join(parts).strip()


def _join(items: Any) -> str:
    if not items:
        return ""
    if isinstance(items, str):
        return items
    out: list[str] = []
    for item in items:
        if isinstance(item, dict):
            out.append(str(item.get("name") or item.get("skill") or item.get("title") or "").strip())
        else:
            out.append(str(item).strip())
    return ", ".join(x for x in out if x)


def _bullets(item: dict) -> list[str]:
    rows = list(item.get("responsibilities") or []) + list(item.get("achievements") or [])
    return [_txt(b) for b in rows if str(b).strip()]


def _write(pdf: FPDF, text: Any, *, size: int = 10, bold: bool = False, italic: bool = False, align: str = "L", color=(35, 35, 35)) -> None:
    body = _txt(text)
    pdf.set_x(pdf.l_margin)
    if not body:
        return
    style = "B" if bold else "I" if italic else ""
    pdf.set_font("Helvetica", style, size)
    pdf.set_text_color(*color)
    pdf.multi_cell(pdf.epw, size * 0.42 + 1.35, body, align=align, new_x="LMARGIN", new_y="NEXT")


def _heading(pdf: FPDF, title: str, color: tuple[int, int, int]) -> None:
    pdf.ln(2.5)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(*color)
    pdf.set_x(pdf.l_margin)
    pdf.cell(pdf.epw, 6, _txt(title.upper()), new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(*color)
    pdf.set_line_width(0.45)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(2.2)
    pdf.set_text_color(35, 35, 35)


def _row(pdf: FPDF, left: str, right: str, *, size: int = 11) -> None:
    pdf.set_x(pdf.l_margin)
    y = pdf.get_y()
    left_w = pdf.epw * 0.68
    pdf.set_font("Helvetica", "B", size)
    pdf.set_text_color(28, 28, 28)
    pdf.cell(left_w, 5.5, _txt(left)[:90], align="L")
    pdf.set_xy(pdf.l_margin + left_w, y)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(90, 90, 90)
    pdf.cell(pdf.epw - left_w, 5.5, _txt(right)[:40], align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(35, 35, 35)


def render_pdf(content: dict[str, Any], template: str = "ats_classic") -> bytes:
    template = template if template in TEMPLATE_LAYOUT else "ats_classic"
    color = TEMPLATE_COLOR[template]
    contact = content.get("contact") or {}
    pdf = ResumePDF(color, sidebar=template == "two_tone")
    if template == "two_tone":
        pdf.set_left_margin(18)
        pdf.set_right_margin(16)
    pdf.add_page()

    name = _txt(contact.get("name") or "Resume")
    meta = "  |  ".join(
        str(x) for x in [contact.get("email"), contact.get("phone"), contact.get("location"), contact.get("links")] if x
    )
    headline = _txt(contact.get("headline") or content.get("target_role") or "")

    if template in ("modern_ats", "executive", "two_tone"):
        bar_h = 36 if template == "executive" else 32
        pdf.set_fill_color(*color)
        pdf.rect(0, 0, pdf.w, bar_h, "F")
        pdf.set_text_color(255, 255, 255)
        pdf.set_xy(pdf.l_margin, 8)
        pdf.set_font("Helvetica", "B", 20 if template == "executive" else 18)
        pdf.cell(pdf.epw, 9, name, align="C" if template == "modern_ats" else "L", new_x="LMARGIN", new_y="NEXT")
        if headline:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_x(pdf.l_margin)
            pdf.cell(pdf.epw, 5, headline, align="C" if template == "modern_ats" else "L", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 8.5)
        pdf.set_x(pdf.l_margin)
        pdf.cell(pdf.epw, 5, meta, align="C" if template == "modern_ats" else "L", new_x="LMARGIN", new_y="NEXT")
        pdf.set_y(bar_h + 8)
        pdf.set_text_color(35, 35, 35)
    else:
        align = "C" if template in ("portfolio",) else "L"
        pdf.set_text_color(*color)
        pdf.set_font("Helvetica", "B", 16 if template == "compact" else 19)
        pdf.set_x(pdf.l_margin)
        pdf.cell(pdf.epw, 9, name, align=align, new_x="LMARGIN", new_y="NEXT")
        if headline:
            _write(pdf, headline, size=10, italic=True, align=align, color=color)
        _write(pdf, meta, size=9, align=align, color=(80, 80, 80))
        pdf.set_draw_color(*color)
        pdf.set_line_width(0.8 if template == "portfolio" else 0.4)
        y = pdf.get_y() + 1
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(4)

    for section in TEMPLATE_LAYOUT[template]:
        if section == "summary" and _txt(content.get("summary")):
            _heading(pdf, "Professional summary" if template == "executive" else "Summary", color)
            _write(pdf, content["summary"], size=10)
        elif section == "skills" and content.get("skills"):
            _heading(pdf, "Technical skills" if template == "technical" else "Skills", color)
            _skills_block(pdf, content["skills"])
        elif section == "experience" and content.get("experience"):
            _heading(pdf, "Leadership & experience" if template == "executive" else "Experience", color)
            _experience_block(pdf, content["experience"])
        elif section == "projects" and content.get("projects"):
            _heading(pdf, "Selected work" if template == "portfolio" else "Projects", color)
            _projects_block(pdf, content["projects"])
        elif section == "education" and content.get("education"):
            _heading(pdf, "Education", color)
            _education_block(pdf, content["education"])

    return bytes(pdf.output())


def _skills_block(pdf: FPDF, skills: Any) -> None:
    if isinstance(skills, dict):
        for k, v in skills.items():
            joined = _join(v) if isinstance(v, list) else str(v or "")
            if not joined:
                continue
            label = str(k).replace("_", " ").title()
            _write(pdf, f"{label}: {joined}", size=9.5)
    else:
        _write(pdf, skills, size=10)


def _experience_block(pdf: FPDF, items: list) -> None:
    for item in items or []:
        if not isinstance(item, dict):
            continue
        title = f"{item.get('title', '')}  ·  {item.get('company', '')}".strip(" ·")
        dates = " – ".join(str(x) for x in [item.get("start_date"), item.get("end_date")] if x)
        _row(pdf, title, dates)
        for bullet in _bullets(item):
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(pdf.epw, 4.8, _txt(f"-  {bullet}"), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)


def _projects_block(pdf: FPDF, items: list) -> None:
    for item in items or []:
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "Project")
        role = str(item.get("role") or "")
        _row(pdf, f"{name}{('  ·  ' + role) if role else ''}", "")
        if item.get("description"):
            _write(pdf, item.get("description"), size=10)
        techs = _join(item.get("technologies") or [])
        if techs:
            _write(pdf, techs, size=9, italic=True, color=(80, 80, 80))
        pdf.ln(1.5)


def _education_block(pdf: FPDF, items: list) -> None:
    for item in items or []:
        if not isinstance(item, dict):
            continue
        left = f"{item.get('degree', '')}  ·  {item.get('institution', '')}".strip(" ·")
        right = str(item.get("graduation_date") or item.get("end_date") or "")
        _row(pdf, left, right)
        extra = "  ·  ".join(str(x) for x in [item.get("major"), f"GPA {item['gpa']}" if item.get("gpa") else ""] if x)
        if extra:
            _write(pdf, extra, size=9, color=(80, 80, 80))
        pdf.ln(1)


def render_docx(content: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    contact = content.get("contact") or {}
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(str(contact.get("name") or "Resume"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(28, 28, 28)

    meta = "  |  ".join(str(x) for x in [contact.get("email"), contact.get("phone"), contact.get("location")] if x)
    if meta:
        p = doc.add_paragraph(meta)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if content.get("summary"):
        h = doc.add_heading("Summary", 1)
        h.runs[0].font.color.rgb = RGBColor(28, 28, 28)
        doc.add_paragraph(str(content["summary"]))

    skills = content.get("skills")
    if skills:
        doc.add_heading("Skills", 1)
        if isinstance(skills, dict):
            for k, v in skills.items():
                joined = _join(v) if isinstance(v, list) else str(v or "")
                if joined:
                    doc.add_paragraph(f"{str(k).replace('_', ' ').title()}: {joined}")
        else:
            doc.add_paragraph(str(skills))

    if content.get("experience"):
        doc.add_heading("Experience", 1)
        for item in content["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{item.get('title', '')}  ·  {item.get('company', '')}")
            run.bold = True
            dates = " – ".join(str(x) for x in [item.get("start_date"), item.get("end_date")] if x)
            if dates:
                r2 = p.add_run(f"    {dates}")
                r2.italic = True
            for bullet in _bullets(item):
                doc.add_paragraph(bullet, style="List Bullet")

    if content.get("projects"):
        doc.add_heading("Projects", 1)
        for item in content["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(str(item.get("name") or ""))
            run.bold = True
            if item.get("description"):
                doc.add_paragraph(str(item["description"]))

    if content.get("education"):
        doc.add_heading("Education", 1)
        for item in content["education"]:
            doc.add_paragraph(f"{item.get('degree', '')}  ·  {item.get('institution', '')}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_markdown(content: dict[str, Any]) -> str:
    contact = content.get("contact") or {}
    lines = [f"# {contact.get('name') or 'Resume'}", ""]
    if content.get("summary"):
        lines += ["## Summary", str(content["summary"]), ""]
    lines.append("## Skills")
    skills = content.get("skills") or {}
    if isinstance(skills, dict):
        for k, v in skills.items():
            joined = _join(v) if isinstance(v, list) else str(v or "")
            if joined:
                lines.append(f"- **{str(k).replace('_', ' ').title()}:** {joined}")
    lines.append("")
    if content.get("experience"):
        lines.append("## Experience")
        for item in content["experience"]:
            lines.append(f"### {item.get('title', '')}  ·  {item.get('company', '')}")
            for b in _bullets(item):
                lines.append(f"- {b}")
            lines.append("")
    return "\n".join(lines)
